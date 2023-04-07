import docx

# get the filename from the user
filename = input("Enter the filename: ")

# create a new document
doc = docx.Document()

# add some text to the document
doc.add_paragraph('Hello, world!')
doc.add_heading('My Heading', level=1)

# save the document to disk with the user-specified filename
doc.save(f"{filename}.docx")