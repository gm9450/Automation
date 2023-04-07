# D:\Myproject\Automation\Automation\makedocx
# D:\Myproject\Automation\carbon
# D:\Myproject\Automation\capture

import os
import docx

# Get the path of the first folder
folder_path1 = input("Enter the path of the first folder: ")

# Get the list of files in the first folder
files1 = os.listdir(folder_path1)

# Filter the list to include only .py files
py_files1 = [file for file in files1 if file.endswith('.py')]

# Create a list of filenames without the .py extension
file_names1 = [file[:-3] for file in py_files1]

# Print the list of filenames without the .py extension in the first folder
print("The filenames without the .py extension in the first folder are:")
print(file_names1)

# Get the path of the second folder
folder_path2 = input("Enter the path of the second folder: ")

# Get the list of files in the second folder with .png extension
files2 = os.listdir(folder_path2)

png_files2 = [file for file in files2 if file.endswith('.png')]

# Print the list of .png files in the second folder
print("The .png files in the second folder are:")
print(png_files2)

# Get the path of the third folder
folder_path3 = input("Enter the path of the third folder: ")

# Get the list of files in the third folder with .png extension
files3 = os.listdir(folder_path3)

png_files3 = [file for file in files3 if file.endswith('.png')]

# Print the list of .png files in the third folder
print("The .png files in the third folder are:")
print(png_files3)

# create a new document
doc = docx.Document()

# add some text and images to the document
for i, file_name in enumerate(file_names1):
    # add the text from the first list
    code_name = input("코드이름 입력 : ")
    doc.add_paragraph(f"{code_name} <{file_name}>")
    
    # add the image from the second list
    if len(png_files2) > i:
        image_path = os.path.join(folder_path2, png_files2[i])
        doc.add_picture(image_path, width=docx.shared.Cm(15.17), height=docx.shared.Cm(21.78))
    
    # add the images from the third list that match the current letter in the first list
    for png_file in png_files3:
        if file_name in png_file:
            image_path = os.path.join(folder_path3, png_file)
            doc.add_picture(image_path, width=docx.shared.Cm(14.88), height=docx.shared.Cm(7.28))
    doc.add_page_break()

# save the document to disk with the user-specified filename
filename = input("Enter the filename: ")
doc.save(f"{filename}.docx")
