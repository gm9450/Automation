import os
import docx
import tkinter as tk
from tkinter import filedialog, messagebox

root = tk.Tk()
root.title("Document Creator")

def get_folder_path():
    folder_path = filedialog.askdirectory()
    return folder_path

def create_doc():
    # Get the paths of the folders
    folder_path1 = folder_path1_entry.get()
    folder_path2 = folder_path2_entry.get()
    folder_path3 = folder_path3_entry.get()

# Get the list of files in the first folder
files1 = os.listdir(folder_path1)

# Filter the list to include only .py files
py_files1 = [file for file in files1 if file.endswith('.py')]

# Create a list of filenames without the .py extension
file_names1 = [file[:-3] for file in py_files1]

# Get the list of files in the second folder with .png extension
files2 = os.listdir(folder_path2)
png_files2 = [file for file in files2 if file.endswith('.png')]

# Get the list of files in the third folder with .png extension
files3 = os.listdir(folder_path3)
png_files3 = [file for file in files3 if file.endswith('.png')]

# Create a new document
doc = docx.Document()

# Add some text and images to the document
for i, file_name in enumerate(file_names1):
    # Add the text from the first list
    code_name = code_name_entry.get()
    doc.add_paragraph(f"{code_name} <{file_name}>")

    # Add the image from the second list
    if len(png_files2) > i:
        image_path = os.path.join(folder_path2, png_files2[i])
        doc.add_picture(image_path, width=docx.shared.Cm(15.17), height=docx.shared.Cm(21.78))

    # Add the images from the third list that match the current letter in the first list
    for png_file in png_files3:
        if file_name in png_file:
            image_path = os.path.join(folder_path3, png_file)
            doc.add_picture(image_path, width=docx.shared.Cm(14.88), height=docx.shared.Cm(7.28))
    doc.add_page_break()

# Save the document to disk with the user-specified filename
filename = filename_entry.get()
doc.save(f"{filename}.docx")

# Show a message box indicating that the document was created
messagebox.showinfo("Document Creator", "Document created successfully.")

folder_path1_label = tk.Label(root, text="First folder path:")
folder_path1_entry = tk.Entry(root)
folder_path1_button = tk.Button(root, text="Browse", command=lambda: folder_path1_entry.insert(0, get_folder_path()))

folder_path2_label = tk.Label(root, text="Second folder path:")
folder_path2_entry = tk.Entry(root)
folder_path2_button = tk.Button(root, text="Browse", command=lambda: folder_path2_entry.insert(0, get_folder_path()))

folder_path3_label = tk.Label(root, text="Third folder path:")
folder_path3_entry = tk.Entry(root)
folder_path3_button = tk.Button(root, text="Browse", command=lambda: folder_path3_entry.insert(0, get_folder_path()))
